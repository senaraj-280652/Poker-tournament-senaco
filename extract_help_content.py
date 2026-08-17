# -*- coding: utf-8 -*-
"""Régénère help_content.json à partir de MANUEL_UTILISATEUR_TOURNOI_CPC.docx
— utilisé par l'aide intégrée de l'application (help_browser.py, menu Aide /
touche F1). À relancer manuellement après toute mise à jour du manuel :

    python3 extract_help_content.py

Nécessite le paquet python-docx (pip install python-docx) — uniquement pour
cette régénération ponctuelle, PAS une dépendance de l'application elle-même
(qui ne lit que le JSON déjà généré, voir help_browser.py).

Parcourt le corps du document dans l'ordre réel (paragraphes ET tableaux
interleaved, contrairement à Document.paragraphs / .tables qui les séparent
et perdent l'ordre) : chaque titre (Heading 1/2/3) démarre une nouvelle
entrée ; le texte normal, les puces et les tableaux rencontrés avant le
titre suivant sont ajoutés à son corps. La page de titre et le Sommaire
(sa propre table des matières) sont ignorés — le navigateur d'aide
reconstruit sa propre arborescence cliquable à partir des niveaux de titre."""
import json

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DOCX = "MANUEL_UTILISATEUR_TOURNOI_CPC.docx"
OUTPUT_JSON = "help_content.json"

HEADING_LEVELS = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}


def iter_block_items(document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def table_to_text(table):
    lines = []
    for i, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        line = " — ".join(cells)
        lines.append(line)
        if i == 0:
            lines.append("-" * min(60, len(line)))
    return "\n".join(lines)


def extract(document):
    entries = []
    started = False
    current = None

    for block in iter_block_items(document):
        if isinstance(block, Table):
            if started and current is not None:
                current["body"].append(table_to_text(block))
            continue

        style = block.style.name
        text = block.text.strip()

        if style == "Heading 1" and text == "Sommaire":
            continue

        if style in HEADING_LEVELS:
            if not text:
                continue  # paragraphes de titre vides (espacement)
            if not started:
                if style != "Heading 1":
                    continue
                started = True
            current = {"level": HEADING_LEVELS[style], "title": text, "body": []}
            entries.append(current)
            continue

        if not started or not text:
            continue

        if style == "List Bullet":
            current["body"].append("• " + text)
        elif style == "Normal":
            current["body"].append(text)

    for e in entries:
        e["body"] = "\n\n".join(e["body"])
    return entries


def main():
    document = docx.Document(SOURCE_DOCX)
    entries = extract(document)
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)
    print(f"{len(entries)} sections écrites dans {OUTPUT_JSON}")


if __name__ == "__main__":
    main()
